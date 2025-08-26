import os
import uuid
import docx
import textract
import google.generativeai as genai
from config.config import Config
from langchain_postgres.vectorstores import PGVector
from langchain.docstore.document import Document
from typing import List, Optional
from PyPDF2 import PdfReader

# Embedding generator using Google Generative AI
class FileEmbeddingGenerator:
    def __init__(self, embedding_model_name="models/embedding-001"):
        self.embedding_model_name = embedding_model_name
        genai.configure(api_key=Config.GOOGLE_API_KEY)
        self.embedding_model = genai.GenerativeModel(self.embedding_model_name)


    def generate_embeddings(self, texts: list[str]) -> list[list[float]]:
        try:
            embeddings = []
            for text in texts:
                response = genai.embed_content(
                    model=self.embedding_model_name,
                    content=text,
                    task_type="retrieval_document",
                    title=text[:50]
                )
                embeddings.append(response['embedding'])
            return embeddings
        except Exception as e:
            raise Exception(f"Error generating embeddings for texts: {e}")
        
    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return self.generate_embeddings(texts)


# Vector store using Google Generative AI

class PGVectorStore:
    def __init__(self):
        """
        Initializes the PGVectorStore class, connecting to PostgreSQL using the connection string in the Config class.
        Uses the provided collection_name to store the embeddings.
        """
        try:
            connection = Config.POSTGRES_CONNECTION
            collection_name = "file_embeddings"
            embedding_generator = FileEmbeddingGenerator()
            self.vector_store = PGVector(
                embeddings=embedding_generator,
                connection=connection,
                collection_name=collection_name,
                use_jsonb=True
            )
        except Exception as e:
            raise Exception(f"Error connecting to PostgreSQL: {e}")


    def store_embeddings(self, chunks: list, embeddings: list):
        try:
            documents = []
            expected_dimension = None

            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                if expected_dimension is None:
                    expected_dimension = len(embedding)

                if len(embedding) != expected_dimension:
                    raise ValueError(
                        f"Warning: Skipping chunk {i} due to inconsistent embedding dimension. Expected {expected_dimension}, got {len(embedding)}."
                    )
                embedding = [float(x) for x in embedding]
                doc_metadata = {
                    "content": chunk,
                    "embedding": embedding,
                    "chunk_sequence": i
                }
                doc = Document(page_content=chunk, metadata=doc_metadata)
                doc.id = str(uuid.uuid4())
                documents.append(doc)

            print(f"Number of documents to attempt adding: {len(documents)}")

            if documents:
                added_ids = self.vector_store.add_documents(documents)
                print(f"Successfully added {len(added_ids)} documents.")
            else:
                print("No documents to add.")
        except Exception as e:
            raise Exception(f"Error storing embeddings: {e}")

def _extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF using PyPDF2. Raises a clear error if PyPDF2 is not installed.
    """
    if PdfReader is None:
        raise RuntimeError("PyPDF2 is required to extract text from PDFs. Install with: pip install PyPDF2")
    reader = PdfReader(file_path)
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a .docx using python-docx.
    """
    if docx is None:
        raise RuntimeError("python-docx is required to extract text from .docx files. Install with: pip install python-docx")
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(paragraphs)


def _extract_text_from_doc(file_path: str) -> str:
    """
    Extract text from a legacy .doc using textract if available.
    """
    if textract is None:
        raise RuntimeError("textract is required to extract text from .doc files. Install with: pip install textract")
    # textract returns bytes
    text = textract.process(file_path)
    try:
        return text.decode("utf-8", errors="ignore")
    except Exception:
        return str(text)


def _extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from a .txt file with robust decoding fallback.
    """
    with open(file_path, "rb") as fh:
        raw = fh.read()
    # try utf-8, then fallback to latin-1 to avoid decode errors
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw.decode(enc)
        except Exception:
            continue
    # final fallback
    return raw.decode("utf-8", errors="ignore")


def _chunk_text(text: str, chunk_size: int = 2000, overlap: int = 200) -> List[str]:
    """
    Split text into chunks with a small overlap.
    """
    if chunk_size <= 0:
        raise ValueError("chunk_size must be > 0")
    chunks: List[str] = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + chunk_size, length)
        chunks.append(text[start:end])
        # advance start but keep overlap
        if end == length:
            break
        start = max(0, end - overlap)
    return chunks


def process_file(file_path: str, chunk_size: Optional[int] = None, overlap: int = 200) -> dict:
    """
    Main entry to process a file:
    - Extract text (PDFs via PyPDF2, others via reading)
    - Chunk text
    - Generate embeddings with FileEmbeddingGenerator
    - Store vectors using PGVectorStore.store_embeddings
    Returns a small summary dict.
    """
    try:
        # try to pick up CHUNK_SIZE if available
        if chunk_size is None:
            try:
                from constant.file_constant import CHUNK_SIZE as DEFAULT_CHUNK
                chunk_size = DEFAULT_CHUNK
            except Exception:
                chunk_size = 2000

        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            text = _extract_text_from_pdf(file_path)
        elif ext == ".docx":
            text = _extract_text_from_docx(file_path)
        elif ext == ".doc":
            text = _extract_text_from_doc(file_path)
        elif ext == ".txt":
            text = _extract_text_from_txt(file_path)
        else:
            # generic read fallback for other text-like files
            with open(file_path, "r", encoding="utf-8", errors="ignore") as fh:
                text = fh.read()

        if not text or not text.strip():
            return {"status": "no_text_extracted", "chunks": 0}

        chunks = _chunk_text(text, chunk_size=int(chunk_size), overlap=int(overlap))
        print(f"[upload_processor] Created {len(chunks)} chunks from {file_path}")

        embedding_gen = FileEmbeddingGenerator()
        embeddings = embedding_gen.embed_documents(chunks)

        pg_store = PGVectorStore()
        pg_store.store_embeddings(chunks, embeddings)
        return {"status": "processed", "chunks": len(chunks)}
    except Exception as err:
        print(f"[upload_processor] Error processing file {file_path}: {err}")
        raise

